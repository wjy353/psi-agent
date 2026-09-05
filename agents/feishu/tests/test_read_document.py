"""Tests for the local DOCX reader used by project learning."""

import importlib.util
import sys
from pathlib import Path

import anyio
import pytest
from docx import Document

TOOLS_DIR = Path(__file__).resolve().parents[1] / "tools"


def _load_tool():
    sys.path.insert(0, str(TOOLS_DIR))
    try:
        spec = importlib.util.spec_from_file_location("haitun_read_document", TOOLS_DIR / "read_document.py")
        assert spec is not None and spec.loader is not None
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        return module
    finally:
        sys.path.pop(0)


@pytest.mark.anyio
async def test_read_document_extracts_paragraphs_tables_and_source_path(tmp_path: Path, monkeypatch) -> None:
    source = tmp_path / "evidence.docx"
    document = Document()
    document.add_heading("项目需求", level=1)
    document.add_paragraph("至少三位差异化专家。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "要求"
    table.cell(1, 0).text = "source_role"
    table.cell(1, 1).text = "按契约定义"
    document.save(str(source))

    tool = _load_tool()
    monkeypatch.setattr(tool._paths, "resolve_user_path", lambda value: anyio.Path(value))
    result = await tool.read_document(str(source))

    assert f"[Source: {source}]" in result
    assert "[Extraction: paragraphs and tables in document order;" in result
    assert "项目需求" in result
    assert "至少三位差异化专家。" in result
    assert "| 字段 | 要求 |" in result
    assert "| source_role | 按契约定义 |" in result
    assert not result.startswith("PK")


@pytest.mark.anyio
async def test_read_document_rejects_generic_text_files(tmp_path: Path, monkeypatch) -> None:
    source = tmp_path / "notes.txt"
    source.write_text("plain", encoding="utf-8")
    tool = _load_tool()
    monkeypatch.setattr(tool._paths, "resolve_user_path", lambda value: anyio.Path(value))

    result = await tool.read_document(str(source))

    assert result.startswith("[Error] Unsupported document type")
