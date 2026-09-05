"""Write-Word tool - create real .docx reports with consistent Chinese fonts.

The recurring "字体不齐" (uneven font) bug comes from writing a .docx whose runs
only set the Latin font (``w:ascii``/``w:hAnsi``) and never the East-Asian font
(``w:eastAsia``); Word then renders CJK glyphs in its default East-Asian face, so
some Chinese characters fall back to a different typeface. This tool builds the
document from structured content and sets ``w:eastAsia`` on every base style once,
so all text — paragraphs, headings, and table cells — is consistent by default.
"""

from __future__ import annotations

import json
from typing import Any

import _runtime_paths as _paths
import anyio
from docx import Document
from docx.oxml.ns import qn


def _set_cjk_font(doc: Any, cjk: str, latin: str) -> None:
    """Set East-Asian + Latin fonts on every base style so all text is consistent.

    Uses ``rPr.get_or_add_rFonts()`` (not a raw append) so the ``w:rFonts``
    element lands as the first child of ``w:rPr`` — the order the OOXML schema
    requires. Iterating ``doc.styles`` covers Normal, every Heading, Title, and
    the table styles in one pass.
    """
    for style in doc.styles:
        element = style.element
        if not hasattr(element, "get_or_add_rPr"):
            continue  # e.g. numbering styles have no run properties
        rpr = element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), latin)
        rfonts.set(qn("w:hAnsi"), latin)
        rfonts.set(qn("w:eastAsia"), cjk)  # the attribute that fixes 字体不齐


def _add_block(doc: Any, block: dict[str, Any]) -> None:
    """Render one content block into the document."""
    kind = block.get("type")
    if kind == "heading":
        level = int(block.get("level", 1))
        doc.add_heading(str(block.get("text", "")), level=max(0, min(level, 9)))
    elif kind == "paragraph":
        doc.add_paragraph(str(block.get("text", "")))
    elif kind == "table":
        rows = block.get("rows") or []
        if not rows:
            return
        ncols = max(len(r) for r in rows)
        table = doc.add_table(rows=0, cols=ncols)
        table.style = block.get("style", "Light Grid Accent 1")
        for row in rows:
            cells = table.add_row().cells
            for idx in range(ncols):
                cells[idx].text = "" if idx >= len(row) else str(row[idx])
    elif kind == "page_break":
        doc.add_page_break()
    else:
        raise ValueError(f"unknown block type: {kind!r}")


def _build_document(
    file_path: str,
    title: str,
    blocks: list[dict[str, Any]],
    cjk_font: str,
    latin_font: str,
) -> int:
    """Build and save a .docx synchronously. Returns the number of blocks written."""
    doc = Document()
    _set_cjk_font(doc, cjk_font, latin_font)  # right after creating the document
    if title:
        doc.add_heading(title, level=0)
    for block in blocks:
        _add_block(doc, block)
    doc.save(file_path)
    return len(blocks)


async def write_word(
    file_path: str,
    blocks_json: list[dict[str, Any]],
    title: str = "",
    cjk_font: str = "微软雅黑",
    latin_font: str = "Calibri",
) -> str:
    """Create a real Word (.docx) report from structured content.

    Use this instead of hand-writing a python-docx script when the user asks for
    a Word document. It sets the East-Asian font (``w:eastAsia``) on every base
    style, so Chinese text renders in one consistent typeface — this is the fix
    for the "字体不齐" (uneven font) bug that appears when only ``run.font.name``
    is set. Drop to raw python-docx only for styling this tool cannot express
    (images, native charts, custom run formatting).

    Args:
        file_path: Output path for the .docx file (e.g. "report.docx").
            Relative paths (including bare filenames) resolve under the
            **session workspace** — never under the agent package directory.
            Absolute paths are used as-is, except writes into the agent
            package itself, which are refused.
        blocks_json: Array of content blocks, in order. Each block
            is an object with a ``type``:
              - ``{"type": "heading", "level": 1, "text": "概述"}`` — level 0 is
                the title style, 1-3 feed the table of contents.
              - ``{"type": "paragraph", "text": "本季度……"}``
              - ``{"type": "table", "rows": [["月份", "收入"], ["1月", "100"]],
                "style": "Light Grid Accent 1"}`` — ``style`` is optional.
              - ``{"type": "page_break"}``
            Example: '[{"type":"heading","level":1,"text":"概述"},
                       {"type":"paragraph","text":"正文"}]'.
        title: Optional document title rendered with the Title style at the top.
        cjk_font: East-Asian font for Chinese text (default 微软雅黑). Safe
            alternatives: 宋体, 黑体.
        latin_font: Latin font for ASCII text (default Calibri).

    Returns:
        Success message with the block count and the **resolved** output path
        (the path a ``[SEND:…]`` marker should point at), or an error message.
    """
    if not file_path.lower().endswith(".docx"):
        file_path = f"{file_path}.docx"

    blocks: Any = blocks_json
    if isinstance(blocks, str):
        try:
            blocks = json.loads(blocks)
        except json.JSONDecodeError as e:
            return f"[Error] blocks_json is not valid JSON: {e}"

    if not isinstance(blocks, list) or not all(isinstance(b, dict) for b in blocks):
        return '[Error] blocks_json must be an array of objects, e.g. [{"type":"paragraph","text":"hi"}]'
    if not blocks and not title:
        return "[Error] provide a title or at least one block"

    # Deliverables belong in the session workspace: bare/relative names resolve
    # there (the framework binds get_workspace() per turn), never against the
    # process cwd — which is the agent package dir. Writing into the package
    # silently breaks [SEND:] delivery (markers point at the workspace) and
    # pollutes the version-controlled tree, so it is refused outright.
    path = _paths.resolve_user_path(file_path)
    refusal = _paths.refuse_agent_write(str(path))
    if refusal is not None:
        return refusal
    parent = path.parent
    if not await parent.exists():
        await parent.mkdir(parents=True, exist_ok=True)

    try:
        count = await anyio.to_thread.run_sync(  # ty: ignore
            _build_document, str(path), title, blocks, cjk_font, latin_font
        )
    except Exception as e:  # python-docx raises assorted errors on bad content
        return f"[Error] Failed to write Word file: {e!r}"

    return f"[OK] Wrote {count} block(s) to {path}"


async def write_word_from_markdown(
    markdown_path: str,
    file_path: str,
    cjk_font: str = "微软雅黑",
    latin_font: str = "Calibri",
) -> str:
    """Convert an existing Markdown artifact into a Word document.

    Prefer this for long contracts, SOPs, and reports: draft the substantial
    content to Markdown first, then call this tool with two short paths. This
    avoids oversized structured tool arguments and runtime package installs.

    Args:
        markdown_path: Existing UTF-8 Markdown source file. Relative paths
            resolve under the session workspace.
        file_path: Output .docx path (see ``write_word`` for resolution rules).
        cjk_font: East-Asian font for Chinese text.
        latin_font: Latin font for ASCII text.

    Returns:
        Success message with the block count, or an error message.
    """
    source = _paths.resolve_user_path(markdown_path)
    if not await source.exists():
        return f"[Error] Markdown source does not exist: {markdown_path}"
    try:
        markdown = await source.read_text(encoding="utf-8")
    except OSError as e:
        return f"[Error] Failed to read Markdown source: {e}"

    blocks: list[dict[str, Any]] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            blocks.append({"type": "paragraph", "text": "\n".join(paragraph_lines)})
            paragraph_lines.clear()

    for line in markdown.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            marker, separator, heading = stripped.partition(" ")
            if separator and set(marker) == {"#"}:
                flush_paragraph()
                blocks.append({"type": "heading", "level": min(len(marker), 9), "text": heading})
                continue
        if not stripped:
            flush_paragraph()
        else:
            paragraph_lines.append(line)
    flush_paragraph()
    if not blocks:
        return "[Error] Markdown source is empty"
    return await write_word(file_path, blocks, cjk_font=cjk_font, latin_font=latin_font)
